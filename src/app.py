import streamlit as st
import os
import pandas as pd
from docx import Document
import re
from pathlib import Path
import plotly.express as px
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from collections import Counter
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns
from categorizer import DocumentCategorizer
import io
from datetime import datetime
import json

# Page configuration
st.set_page_config(
    page_title="Document Search Dashboard",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

class DocumentSearcher:
    def __init__(self):
        self.documents = []
        self.vectorizer = None
        self.tfidf_matrix = None
        self.categorizer = DocumentCategorizer()
        
    def load_documents(self, folder_path):
        """Load all Word documents from the specified folder"""
        self.documents = []
        doc_files = list(Path(folder_path).glob("*.docx"))
        
        if not doc_files:
            st.warning(f"No .docx files found in {folder_path}")
            return
            
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, doc_path in enumerate(doc_files):
            try:
                status_text.text(f"Loading {doc_path.name}...")
                doc = Document(doc_path)
                
                # Extract text from all paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                content = "\n".join(full_text)
                
                # Get basic document info
                doc_info = {
                    'filename': doc_path.name,
                    'path': str(doc_path),
                    'content': content,
                    'word_count': len(content.split()),
                    'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                    'size_kb': round(doc_path.stat().st_size / 1024, 2)
                }
                
                self.documents.append(doc_info)
                
            except Exception as e:
                st.error(f"Error loading {doc_path.name}: {str(e)}")
            
            progress_bar.progress((i + 1) / len(doc_files))
        
        status_text.text("Building search index...")
        self._build_search_index()
        status_text.text("Categorizing documents...")
        self._categorize_documents()
        status_text.text(f"Loaded {len(self.documents)} documents successfully!")
        
    def _build_search_index(self):
        """Build TF-IDF index for similarity search"""
        if not self.documents:
            return
            
        corpus = [doc['content'] for doc in self.documents]
        self.vectorizer = TfidfVectorizer(
            max_features=1000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)
    
    def search(self, query, method='keyword', top_k=10):
        """Search documents using different methods"""
        if not self.documents:
            return []
            
        if method == 'keyword':
            return self._keyword_search(query, top_k)
        elif method == 'similarity':
            return self._similarity_search(query, top_k)
        else:
            return self._combined_search(query, top_k)
    
    def search_by_tags(self, selected_tags, top_k=10):
        """Search documents by selected tags"""
        if not selected_tags or not self.documents:
            return []
            
        results = []
        for doc in self.documents:
            doc_tags = doc.get('tags', [])
            matching_tags = [tag for tag in selected_tags if tag in doc_tags]
            
            if matching_tags:
                score = len(matching_tags) / len(selected_tags)  # Relevance based on tag match ratio
                results.append({
                    **doc,
                    'score': score,
                    'matching_tags': matching_tags,
                    'context': [f"Document tagged with: {', '.join(matching_tags)}"]
                })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]
    
    def _keyword_search(self, query, top_k):
        """Simple keyword-based search"""
        results = []
        query_words = query.lower().split()
        
        for doc in self.documents:
            content_lower = doc['content'].lower()
            score = 0
            matches = []
            
            for word in query_words:
                count = content_lower.count(word)
                score += count
                if count > 0:
                    matches.append(word)
            
            if score > 0:
                # Find context around matches
                context = self._get_context(doc['content'], query_words)
                results.append({
                    **doc,
                    'score': score,
                    'matches': matches,
                    'context': context
                })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]
    
    def _similarity_search(self, query, top_k):
        """TF-IDF based similarity search"""
        if self.vectorizer is None:
            return self._keyword_search(query, top_k)
            
        query_vector = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vector, self.tfidf_matrix).flatten()
        
        results = []
        for i, doc in enumerate(self.documents):
            if similarities[i] > 0:
                context = self._get_context(doc['content'], query.split())
                results.append({
                    **doc,
                    'score': similarities[i],
                    'context': context
                })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]
    
    def _combined_search(self, query, top_k):
        """Combine keyword and similarity search"""
        keyword_results = self._keyword_search(query, top_k * 2)
        similarity_results = self._similarity_search(query, top_k * 2)
        
        # Combine and deduplicate
        combined = {}
        for result in keyword_results + similarity_results:
            filename = result['filename']
            if filename in combined:
                combined[filename]['score'] += result['score']
            else:
                combined[filename] = result
        
        return sorted(combined.values(), key=lambda x: x['score'], reverse=True)[:top_k]
    
    def _get_context(self, content, query_words, context_length=200):
        """Extract context around query matches"""
        content_lower = content.lower()
        contexts = []
        
        for word in query_words:
            word = word.lower()
            start = content_lower.find(word)
            if start != -1:
                context_start = max(0, start - context_length // 2)
                context_end = min(len(content), start + len(word) + context_length // 2)
                context = content[context_start:context_end]
                if context_start > 0:
                    context = "..." + context
                if context_end < len(content):
                    context = context + "..."
                contexts.append(context)
        
        return contexts[:3]  # Return up to 3 contexts

    def _categorize_documents(self):
        """Categorize and tag all documents"""
        if not self.documents:
            return
            
        categorized_docs = []
        for doc in self.documents:
            analysis = self.categorizer.analyze_document(
                doc['content'], 
                doc['filename']
            )
            # Merge original document data with analysis
            categorized_doc = {**doc, **analysis}
            categorized_docs.append(categorized_doc)
        
        self.documents = categorized_docs

    def get_category_summary(self):
        """Get summary statistics for categories and tags"""
        return self.categorizer.get_category_summary(self.documents)

    def filter_documents(self, category=None, doc_type=None, tags=None, 
                        urgency=None, min_readability=None, max_readability=None):
        """Filter documents based on various criteria"""
        filtered = self.documents.copy()
        
        if category and category != 'All':
            filtered = [doc for doc in filtered 
                       if doc.get('business_category') == category]
        
        if doc_type and doc_type != 'All':
            filtered = [doc for doc in filtered 
                       if doc.get('document_type') == doc_type]
        
        if tags:
            for tag in tags:
                filtered = [doc for doc in filtered 
                           if tag in doc.get('tags', [])]
        
        if urgency and urgency != 'All':
            filtered = [doc for doc in filtered 
                       if doc.get('urgency_level') == urgency]
        
        if min_readability is not None:
            filtered = [doc for doc in filtered 
                       if doc.get('readability', 0) >= min_readability]
        
        if max_readability is not None:
            filtered = [doc for doc in filtered 
                       if doc.get('readability', 100) <= max_readability]
        
        return filtered

    def generate_wordcloud_tags(self, documents=None):
        """Generate word cloud for document tags"""
        if documents is None:
            documents = self.documents
        
        if not documents:
            return None
            
        # Collect all tags with frequencies
        all_tags = []
        for doc in documents:
            all_tags.extend(doc.get('tags', []))
        
        if not all_tags:
            return None
            
        tag_freq = Counter(all_tags)
        
        # Generate word cloud
        wordcloud = WordCloud(
            width=800, 
            height=400, 
            background_color='white',
            max_words=100,
            colormap='viridis',
            relative_scaling=0.5,
            random_state=42
        ).generate_from_frequencies(tag_freq)
        
        return wordcloud, tag_freq

    def create_summary_report(self, documents, search_query="", search_method="", 
                            include_details=True, include_analytics=True, 
                            include_tag_cloud=True):
        """Create a comprehensive summary report"""
        if not documents:
            return None
            
        report_data = {
            'metadata': {
                'generated_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'total_documents': len(documents),
                'search_query': search_query,
                'search_method': search_method
            },
            'summary_statistics': {
                'total_word_count': sum(doc.get('word_count', 0) for doc in documents),
                'average_readability': round(np.mean([doc.get('readability', 0) for doc in documents]), 1),
                'total_size_kb': round(sum(doc.get('size_kb', 0) for doc in documents), 2),
                'document_categories': dict(Counter([doc.get('business_category', 'Unknown') for doc in documents])),
                'document_types': dict(Counter([doc.get('document_type', 'Unknown') for doc in documents])),
                'urgency_levels': dict(Counter([doc.get('urgency_level', 'Unknown') for doc in documents])),
                'sentiment_distribution': dict(Counter([doc.get('sentiment', {}).get('label', 'Unknown') for doc in documents]))
            }
        }
        
        if include_details:
            report_data['document_details'] = []
            for doc in documents:
                doc_summary = {
                    'filename': doc.get('filename', ''),
                    'business_category': doc.get('business_category', ''),
                    'document_type': doc.get('document_type', ''),
                    'word_count': doc.get('word_count', 0),
                    'readability': doc.get('readability', 0),
                    'urgency_level': doc.get('urgency_level', ''),
                    'sentiment': doc.get('sentiment', {}),
                    'tags': doc.get('tags', [])[:10],  # Top 10 tags
                    'key_topics': doc.get('key_topics', [])[:5],  # Top 5 topics
                    'entities': doc.get('entities', {})
                }
                report_data['document_details'].append(doc_summary)
        
        if include_analytics:
            all_tags = [tag for doc in documents for tag in doc.get('tags', [])]
            report_data['analytics'] = {
                'top_tags': dict(Counter(all_tags).most_common(20)),
                'readability_stats': {
                    'min': min([doc.get('readability', 0) for doc in documents]),
                    'max': max([doc.get('readability', 0) for doc in documents]),
                    'median': np.median([doc.get('readability', 0) for doc in documents])
                },
                'complexity_stats': {
                    'min': min([doc.get('complexity_score', 0) for doc in documents]),
                    'max': max([doc.get('complexity_score', 0) for doc in documents]),
                    'average': round(np.mean([doc.get('complexity_score', 0) for doc in documents]), 1)
                }
            }
        
        return report_data

# Initialize the searcher
@st.cache_resource
def get_searcher():
    return DocumentSearcher()

def create_downloadable_summary(report_data, format_type="json"):
    """Create downloadable summary in different formats"""
    if format_type == "json":
        return json.dumps(report_data, indent=2, ensure_ascii=False)
    
    elif format_type == "csv":
        # Create CSV for document details
        if 'document_details' in report_data:
            df = pd.DataFrame(report_data['document_details'])
            return df.to_csv(index=False)
    
    elif format_type == "text":
        # Create human-readable text summary
        text_summary = f"""
DOCUMENT SEARCH SUMMARY REPORT
Generated: {report_data['metadata']['generated_date']}

OVERVIEW
--------
Total Documents: {report_data['metadata']['total_documents']}
Search Query: {report_data['metadata']['search_query'] or 'N/A'}
Search Method: {report_data['metadata']['search_method'] or 'N/A'}

STATISTICS
----------
Total Word Count: {report_data['summary_statistics']['total_word_count']:,}
Average Readability: {report_data['summary_statistics']['average_readability']}
Total Size: {report_data['summary_statistics']['total_size_kb']} KB

DOCUMENT CATEGORIES
-------------------
"""
        for category, count in report_data['summary_statistics']['document_categories'].items():
            text_summary += f"{category}: {count}\n"
        
        text_summary += "\nDOCUMENT TYPES\n--------------\n"
        for doc_type, count in report_data['summary_statistics']['document_types'].items():
            text_summary += f"{doc_type}: {count}\n"
        
        if 'analytics' in report_data:
            text_summary += "\nTOP TAGS\n--------\n"
            for tag, count in list(report_data['analytics']['top_tags'].items())[:10]:
                text_summary += f"{tag}: {count}\n"
        
        return text_summary
    
    return None

def main():
    st.title("📄 Document Search Dashboard")
    st.markdown("Search through your Word documents with ease!")
    
    searcher = get_searcher()
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Folder selection
        doc_folder = st.text_input(
            "Documents Folder Path",
            value="./documents",
            help="Path to folder containing .docx files"
        )
        
        if st.button("🔄 Load Documents", type="primary"):
            if os.path.exists(doc_folder):
                with st.spinner("Loading documents..."):
                    searcher.load_documents(doc_folder)
                st.rerun()
            else:
                st.error("Folder not found!")
        
        # Filters section
        if searcher.documents:
            st.header("🏷️ Filters")
            
            # Get available filter options
            categories = list(set([doc.get('business_category', 'General') 
                                 for doc in searcher.documents]))
            doc_types = list(set([doc.get('document_type', 'Unknown') 
                                for doc in searcher.documents]))
            all_tags = list(set([tag for doc in searcher.documents 
                               for tag in doc.get('tags', [])]))
            
            # Category filter
            selected_category = st.selectbox(
                "Business Category",
                ['All'] + sorted(categories),
                key="category_filter"
            )
            
            # Document type filter
            selected_doc_type = st.selectbox(
                "Document Type",
                ['All'] + sorted(doc_types),
                key="doc_type_filter"
            )
            
            # Tags filter
            selected_tags = st.multiselect(
                "Tags",
                sorted(all_tags)[:20],  # Limit to top 20 tags
                key="tags_filter"
            )
            
            # Urgency filter
            selected_urgency = st.selectbox(
                "Urgency Level",
                ['All', 'High', 'Medium', 'Low'],
                key="urgency_filter"
            )
            
            # Readability filter
            readability_range = st.slider(
                "Readability Score Range",
                min_value=0,
                max_value=100,
                value=(0, 100),
                key="readability_filter",
                help="Flesch Reading Ease score (higher = easier to read)"
            )
        
        # Search settings
        st.header("🔍 Search Settings")
        search_method = st.selectbox(
            "Search Method",
            ["keyword", "similarity", "combined", "tags"],
            help="Choose how to search through documents"
        )
        
        max_results = st.slider("Max Results", 1, 50, 10)
    
    # Main content area
    if not searcher.documents:
        st.info("👈 Please load documents using the sidebar to get started!")
        
        # Show sample folder structure
        st.markdown("### Expected Folder Structure")
        st.code("""
documents/
├── document1.docx
├── document2.docx
├── document3.docx
└── ...
        """)
    else:
        # Document statistics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📁 Total Documents", len(searcher.documents))
        with col2:
            total_words = sum(doc['word_count'] for doc in searcher.documents)
            st.metric("📝 Total Words", f"{total_words:,}")
        with col3:
            total_size = sum(doc['size_kb'] for doc in searcher.documents)
            st.metric("💾 Total Size", f"{total_size:.1f} KB")
        with col4:
            avg_words = total_words // len(searcher.documents) if searcher.documents else 0
            st.metric("📊 Avg Words/Doc", f"{avg_words:,}")
        
        # Apply filters to get filtered documents
        filtered_docs = searcher.documents.copy()
        if searcher.documents:
            # Apply filters
            category_filter = selected_category if selected_category != 'All' else None
            doc_type_filter = selected_doc_type if selected_doc_type != 'All' else None
            urgency_filter = selected_urgency if selected_urgency != 'All' else None
            
            filtered_docs = searcher.filter_documents(
                category=category_filter,
                doc_type=doc_type_filter,
                tags=selected_tags,
                urgency=urgency_filter,
                min_readability=readability_range[0],
                max_readability=readability_range[1]
            )
            
            # Update searcher's documents temporarily for search
            original_docs = searcher.documents
            searcher.documents = filtered_docs
            searcher._build_search_index()

        # Category Overview Dashboard
        if searcher.documents:
            st.markdown("---")
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Overview", "🔍 Search", "☁️ Tag Cloud Search", "📈 Analytics", "📄 Summary Report"])
            
            with tab1:
                st.header("📊 Document Overview")
                summary = searcher.get_category_summary()
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("📁 Total Documents", len(filtered_docs))
                with col2:
                    total_words = sum(doc.get('word_count', 0) for doc in filtered_docs)
                    st.metric("📝 Total Words", f"{total_words:,}")
                with col3:
                    avg_readability = round(np.mean([doc.get('readability', 0) 
                                                   for doc in filtered_docs]), 1)
                    st.metric("📖 Avg Readability", f"{avg_readability}")
                with col4:
                    if summary.get('sentiment_distribution'):
                        most_common_sentiment = max(summary['sentiment_distribution'], 
                                                  key=summary['sentiment_distribution'].get)
                        st.metric("😊 Dominant Sentiment", most_common_sentiment)
                
                # Category and Type Distribution
                col1, col2 = st.columns(2)
                with col1:
                    if len(filtered_docs) > 0:
                        categories = [doc.get('business_category', 'General') 
                                    for doc in filtered_docs]
                        category_counts = pd.DataFrame({
                            'Category': list(Counter(categories).keys()),
                            'Count': list(Counter(categories).values())
                        })
                        fig = px.pie(category_counts, values='Count', names='Category',
                                   title='Business Categories')
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    if len(filtered_docs) > 0:
                        doc_types = [doc.get('document_type', 'Unknown') 
                                   for doc in filtered_docs]
                        type_counts = pd.DataFrame({
                            'Type': list(Counter(doc_types).keys()),
                            'Count': list(Counter(doc_types).values())
                        })
                        fig = px.bar(type_counts, x='Type', y='Count',
                                   title='Document Types')
                        fig.update_layout(xaxis_tickangle=45)
                        st.plotly_chart(fig, use_container_width=True)
                
                # Tags Word Cloud
                if filtered_docs:
                    all_tags = [tag for doc in filtered_docs for tag in doc.get('tags', [])]
                    if all_tags:
                        st.subheader("🏷️ Top Tags")
                        tag_counts = Counter(all_tags)
                        top_tags = pd.DataFrame({
                            'Tag': list(tag_counts.keys())[:20],
                            'Frequency': list(tag_counts.values())[:20]
                        })
                        fig = px.bar(top_tags, x='Frequency', y='Tag', 
                                   orientation='h', title='Most Common Tags')
                        st.plotly_chart(fig, use_container_width=True)
            
            with tab2:
                # Search interface
                search_query = st.text_input(
                    "🔍 Search Query",
                    placeholder="Enter keywords to search for...",
                    help="Search through filtered documents"
                )
                
                if search_query:
                    with st.spinner("Searching..."):
                        if search_method == "tags":
                            # Tag-based search
                            query_tags = [tag.strip() for tag in search_query.split(',')]
                            results = searcher.search_by_tags(query_tags, max_results)
                        else:
                            results = searcher.search(search_query, search_method, max_results)
                    
                    if results:
                        st.success(f"Found {len(results)} results for '{search_query}'")
                        
                        # Display results with enhanced information
                        for i, result in enumerate(results, 1):
                            score_display = f"Score: {result['score']:.3f}"
                            with st.expander(f"📄 {result['filename']} ({score_display})"):
                                col1, col2 = st.columns([2, 1])
                                
                                with col1:
                                    st.markdown(f"**📁 File:** {result['filename']}")
                                    st.markdown(f"**🏢 Category:** {result.get('business_category', 'Unknown')}")
                                    st.markdown(f"**📋 Type:** {result.get('document_type', 'Unknown')}")
                                    st.markdown(f"**📝 Words:** {result['word_count']:,}")
                                    st.markdown(f"**📖 Readability:** {result.get('readability', 0)}")
                                    st.markdown(f"**⚡ Urgency:** {result.get('urgency_level', 'Unknown')}")
                                    
                                    # Show tags
                                    if result.get('tags'):
                                        tags_display = ", ".join(result['tags'][:10])
                                        st.markdown(f"**🏷️ Tags:** {tags_display}")
                                    
                                    # Show matching tags for tag search
                                    if search_method == "tags" and 'matching_tags' in result:
                                        matching_display = ", ".join(result['matching_tags'])
                                        st.markdown(f"**✅ Matching Tags:** {matching_display}")
                                    
                                    # Show sentiment
                                    if result.get('sentiment'):
                                        sentiment = result['sentiment']
                                        st.markdown(f"**😊 Sentiment:** {sentiment['label']} ({sentiment['polarity']:.2f})")
                                    
                                    # Show context if available
                                    if 'context' in result and result['context']:
                                        st.markdown("**🎯 Context:**")
                                        for ctx in result['context']:
                                            # Highlight query terms
                                            highlighted = ctx
                                            for word in search_query.split():
                                                highlighted = re.sub(
                                                    f'({re.escape(word)})',
                                                    r'**\1**',
                                                    highlighted,
                                                    flags=re.IGNORECASE
                                                )
                                            st.markdown(f"*{highlighted}*")
                                
                                with col2:
                                    st.markdown(f"**📊 Match Score:** {result['score']:.3f}")
                                    if 'matches' in result:
                                        matches_display = ", ".join(result['matches'])
                                        st.markdown(f"**🔍 Matched Terms:** {matches_display}")
                                    
                                    # Key topics
                                    if result.get('key_topics'):
                                        st.markdown("**🎯 Key Topics:**")
                                        for topic in result['key_topics'][:3]:
                                            st.markdown(f"• {topic['topic']} ({topic['frequency']})")
                                    
                                    # Full content preview
                                    if st.button(f"👁️ Preview Full Content", key=f"preview_{i}"):
                                        content_preview = result['content'][:2000]
                                        if len(result['content']) > 2000:
                                            content_preview += "..."
                                        st.text_area(
                                            "Full Document Content",
                                            content_preview,
                                            height=300,
                                            key=f"content_{i}"
                                        )
                    else:
                        st.warning(f"No results found for '{search_query}'")
                
                # Show filtered document count
                if len(filtered_docs) < len(searcher.documents if 'original_docs' not in locals() else original_docs):
                    total_docs = len(searcher.documents if 'original_docs' not in locals() else original_docs)
                    st.info(f"Showing {len(filtered_docs)} of {total_docs} documents (filtered)")
            
            with tab3:
                st.header("☁️ Tag Cloud Search")
                st.markdown("Click on tags in the word cloud or select from the list to search documents!")
                
                # Generate word cloud for tags
                wordcloud_data = searcher.generate_wordcloud_tags(filtered_docs)
                
                if wordcloud_data:
                    wordcloud, tag_freq = wordcloud_data
                    
                    # Display word cloud
                    fig, ax = plt.subplots(figsize=(12, 6))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis('off')
                    st.pyplot(fig)
                    plt.close()
                    
                    # Interactive tag selection
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.subheader("🏷️ Select Tags for Search")
                        
                        # Create columns for tag buttons
                        top_tags = list(tag_freq.most_common(30))
                        
                        # Initialize session state for selected tags
                        if 'selected_wordcloud_tags' not in st.session_state:
                            st.session_state.selected_wordcloud_tags = []
                        
                        # Tag selection interface
                        cols = st.columns(5)
                        for i, (tag, freq) in enumerate(top_tags):
                            with cols[i % 5]:
                                if st.button(f"{tag} ({freq})", key=f"tag_btn_{i}", 
                                           type="primary" if tag in st.session_state.selected_wordcloud_tags else "secondary"):
                                    if tag in st.session_state.selected_wordcloud_tags:
                                        st.session_state.selected_wordcloud_tags.remove(tag)
                                    else:
                                        st.session_state.selected_wordcloud_tags.append(tag)
                                    st.rerun()
                        
                        # Show selected tags
                        if st.session_state.selected_wordcloud_tags:
                            st.write("**Selected Tags:**", ", ".join(st.session_state.selected_wordcloud_tags))
                            
                            # Clear selection button
                            if st.button("🗑️ Clear Selection"):
                                st.session_state.selected_wordcloud_tags = []
                                st.rerun()
                            
                            # Search by selected tags
                            with st.spinner("Searching by tags..."):
                                tag_results = searcher.search_by_tags(st.session_state.selected_wordcloud_tags, max_results)
                            
                            if tag_results:
                                st.success(f"Found {len(tag_results)} documents with selected tags")
                                
                                for i, result in enumerate(tag_results, 1):
                                    with st.expander(f"📄 {result['filename']} (Score: {result['score']:.3f})"):
                                        col_a, col_b = st.columns([3, 1])
                                        
                                        with col_a:
                                            st.markdown(f"**📁 File:** {result['filename']}")
                                            st.markdown(f"**🏢 Category:** {result.get('business_category', 'Unknown')}")
                                            st.markdown(f"**📋 Type:** {result.get('document_type', 'Unknown')}")
                                            st.markdown(f"**✅ Matching Tags:** {', '.join(result.get('matching_tags', []))}")
                                            
                                            if result.get('tags'):
                                                all_tags_display = ", ".join(result['tags'][:15])
                                                st.markdown(f"**🏷️ All Tags:** {all_tags_display}")
                                        
                                        with col_b:
                                            st.markdown(f"**📊 Tag Match:** {result['score']:.1%}")
                                            st.markdown(f"**📝 Words:** {result.get('word_count', 0):,}")
                                            st.markdown(f"**📖 Readability:** {result.get('readability', 0)}")
                            else:
                                st.warning("No documents found with the selected tags")
                    
                    with col2:
                        st.subheader("📊 Tag Statistics")
                        
                        # Show tag frequency chart
                        top_10_tags = pd.DataFrame({
                            'Tag': [tag for tag, _ in tag_freq.most_common(10)],
                            'Frequency': [freq for _, freq in tag_freq.most_common(10)]
                        })
                        
                        fig = px.bar(top_10_tags, x='Frequency', y='Tag', orientation='h',
                                   title='Top 10 Most Frequent Tags')
                        fig.update_layout(height=400)
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Tag coverage statistics
                        st.markdown("**Tag Coverage:**")
                        st.markdown(f"• Total unique tags: {len(tag_freq)}")
                        st.markdown(f"• Documents with tags: {len([d for d in filtered_docs if d.get('tags')])}")
                        st.markdown(f"• Avg tags per document: {len([tag for doc in filtered_docs for tag in doc.get('tags', [])]) / len(filtered_docs):.1f}")
                
                else:
                    st.info("No tags found in the current document set. Tags are generated automatically when documents are loaded.")
            
            with tab4:
                st.header("📈 Document Analytics")
                
                if filtered_docs:
                    # Readability analysis
                    col1, col2 = st.columns(2)
                    with col1:
                        readability_scores = [doc.get('readability', 0) for doc in filtered_docs]
                        if readability_scores:
                            fig = px.histogram(
                                x=readability_scores,
                                nbins=20,
                                title='Readability Score Distribution',
                                labels={'x': 'Readability Score', 'y': 'Count'}
                            )
                            st.plotly_chart(fig, use_container_width=True)
                    
                    with col2:
                        # Sentiment analysis
                        sentiments = [doc.get('sentiment', {}).get('label', 'Unknown') 
                                    for doc in filtered_docs]
                        sentiment_counts = pd.DataFrame({
                            'Sentiment': list(Counter(sentiments).keys()),
                            'Count': list(Counter(sentiments).values())
                        })
                        fig = px.pie(sentiment_counts, values='Count', names='Sentiment',
                                   title='Sentiment Distribution')
                        st.plotly_chart(fig, use_container_width=True)
                    
                    # Document complexity vs readability
                    complexity_data = []
                    for doc in filtered_docs:
                        complexity_data.append({
                            'Filename': doc.get('filename', ''),
                            'Readability': doc.get('readability', 0),
                            'Complexity': doc.get('complexity_score', 0),
                            'Word Count': doc.get('word_count', 0),
                            'Category': doc.get('business_category', 'Unknown')
                        })
                    
                    if complexity_data:
                        df_complexity = pd.DataFrame(complexity_data)
                        fig = px.scatter(
                            df_complexity,
                            x='Readability',
                            y='Complexity',
                            size='Word Count',
                            color='Category',
                            hover_name='Filename',
                            title='Document Complexity vs Readability',
                            labels={
                                'Readability': 'Readability Score (higher = easier)',
                                'Complexity': 'Complexity Score'
                            }
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    
                    # Top entities across documents
                    all_entities = {}
                    for doc in filtered_docs:
                        entities = doc.get('entities', {})
                        for entity_type, entity_list in entities.items():
                            if entity_type not in all_entities:
                                all_entities[entity_type] = []
                            all_entities[entity_type].extend(entity_list)
                    
                    if all_entities:
                        st.subheader("🔍 Extracted Entities")
                        entity_cols = st.columns(len(all_entities))
                        for i, (entity_type, entities) in enumerate(all_entities.items()):
                            with entity_cols[i % len(entity_cols)]:
                                st.write(f"**{entity_type.title()}:**")
                                unique_entities = list(set(entities))[:5]
                                for entity in unique_entities:
                                    st.write(f"• {entity}")
            
            with tab5:
                st.header("📄 Summary Report Generator")
                st.markdown("Create a customized, downloadable summary of your document analysis!")
                
                # Report configuration
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.subheader("🔧 Report Configuration")
                    
                    # Report options
                    include_details = st.checkbox("Include detailed document information", value=True)
                    include_analytics = st.checkbox("Include analytics and statistics", value=True)
                    include_tag_cloud = st.checkbox("Include tag analysis", value=True)
                    
                    # Search query for report context
                    report_query = st.text_input(
                        "Search Query (for report context)",
                        placeholder="Optional: Enter the search query this report relates to"
                    )
                    
                    # Report format
                    report_format = st.selectbox(
                        "Report Format",
                        ["json", "text", "csv"],
                        help="Choose the format for your downloadable report"
                    )
                
                with col2:
                    st.subheader("📊 Report Preview")
                    st.metric("Documents in Report", len(filtered_docs))
                    st.metric("Total Words", f"{sum(doc.get('word_count', 0) for doc in filtered_docs):,}")
                    st.metric("Categories", len(set([doc.get('business_category', 'Unknown') for doc in filtered_docs])))
                    st.metric("Unique Tags", len(set([tag for doc in filtered_docs for tag in doc.get('tags', [])])))
                
                # Generate report
                if st.button("📊 Generate Report", type="primary"):
                    with st.spinner("Generating comprehensive report..."):
                        report_data = searcher.create_summary_report(
                            documents=filtered_docs,
                            search_query=report_query,
                            search_method=search_method,
                            include_details=include_details,
                            include_analytics=include_analytics,
                            include_tag_cloud=include_tag_cloud
                        )
                        
                        if report_data:
                            # Create downloadable content
                            downloadable_content = create_downloadable_summary(report_data, report_format)
                            
                            if downloadable_content:
                                # Show preview
                                st.subheader("📋 Report Preview")
                                if report_format == "json":
                                    st.json(report_data['summary_statistics'])
                                elif report_format == "text":
                                    st.text_area("Report Preview", downloadable_content[:2000], height=300)
                                elif report_format == "csv" and 'document_details' in report_data:
                                    st.dataframe(pd.DataFrame(report_data['document_details']).head(10))
                                
                                # Download button
                                filename = f"document_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{report_format}"
                                
                                st.download_button(
                                    label=f"📥 Download {report_format.upper()} Report",
                                    data=downloadable_content,
                                    file_name=filename,
                                    mime={
                                        'json': 'application/json',
                                        'text': 'text/plain',
                                        'csv': 'text/csv'
                                    }[report_format]
                                )
                                
                                st.success(f"Report generated successfully! Click the button above to download your {report_format.upper()} report.")
                            else:
                                st.error("Failed to generate downloadable content. Please try a different format.")
                        else:
                            st.error("Failed to generate report. Please ensure documents are loaded.")
                
                # Report template information
                with st.expander("ℹ️ Report Format Information"):
                    st.markdown("""
                    **JSON Format**: Complete structured data including all metadata, statistics, and document details. 
                    Best for programmatic use or importing into other tools.
                    
                    **Text Format**: Human-readable summary with key statistics and insights. 
                    Perfect for sharing with stakeholders or including in presentations.
                    
                    **CSV Format**: Tabular data of document details that can be opened in Excel or other spreadsheet applications. 
                    Ideal for further data analysis.
                    """)
        
        # Restore original documents
        if 'original_docs' in locals():
            searcher.documents = original_docs
            searcher._build_search_index()

if __name__ == "__main__":
    main()