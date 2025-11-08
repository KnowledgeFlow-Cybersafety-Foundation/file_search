"""
Core document searching functionality
"""

import streamlit as st
from docx import Document
from pathlib import Path
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from collections import Counter
from wordcloud import WordCloud
from typing import List, Dict, Any, Optional, Tuple

import sys
sys.path.append('/Users/colinwork/Documents/GitHub/docker_file_search/src')

from categorizer import DocumentCategorizer
from utils.constants import TFIDF_CONFIG, WORDCLOUD_CONFIG
from utils.constants import CONTEXT_LENGTH, MAX_CONTEXTS


class DocumentSearcher:
    """Main class for document search functionality"""
    
    def __init__(self):
        self.documents = []
        self.vectorizer = None
        self.tfidf_matrix = None
        self.categorizer = DocumentCategorizer()
        
    def load_documents(self, folder_path: str) -> None:
        """Load all Word documents from the specified folder"""
        self.documents = []
        doc_files = list(Path(folder_path).glob("*.docx"))
        
        if not doc_files:
            st.warning(f"No .docx files found in {folder_path}")
            return
            
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, doc_path in enumerate(doc_files):
            try:
                status_text.text(f"Loading {doc_path.name}...")
                doc = Document(doc_path)
                
                # Extract text from all paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                content = "\n".join(full_text)
                
                # Get basic document info
                doc_info = {
                    'filename': doc_path.name,
                    'path': str(doc_path),
                    'content': content,
                    'word_count': len(content.split()),
                    'paragraph_count': len([
                        p for p in doc.paragraphs if p.text.strip()
                    ]),
                    'size_kb': round(doc_path.stat().st_size / 1024, 2)
                }
                
                self.documents.append(doc_info)
                
            except Exception as e:
                st.error(f"Error loading {doc_path.name}: {str(e)}")
            
            progress_bar.progress((i + 1) / len(doc_files))
        
        status_text.text("Building search index...")
        self._build_search_index()
        status_text.text("Categorizing documents...")
        self._categorize_documents()
        status_text.text(
            f"Loaded {len(self.documents)} documents successfully!"
        )
        
    def _build_search_index(self) -> None:
        """Build TF-IDF index for similarity search"""
        if not self.documents:
            return
            
        corpus = [doc['content'] for doc in self.documents]
        self.vectorizer = TfidfVectorizer(**TFIDF_CONFIG)
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)
    
    def search(
        self,
        query: str,
        method: str = 'keyword',
        top_k: int = 10
    ) -> List[Dict[str, Any]]:
        """Search documents using different methods"""
        if not self.documents:
            return []
            
        if method == 'keyword':
            return self._keyword_search(query, top_k)
        elif method == 'similarity':
            return self._similarity_search(query, top_k)
        else:
            return self._combined_search(query, top_k)
    
    def search_by_tags(
        self,
        selected_tags: List[str],
        top_k: int = 10
    ) -> List[Dict[str, Any]]:
        """Search documents by selected tags"""
        if not selected_tags or not self.documents:
            return []
            
        results = []
        for doc in self.documents:
            doc_tags = doc.get('tags', [])
            matching_tags = [tag for tag in selected_tags if tag in doc_tags]
            
            if matching_tags:
                score = len(matching_tags) / len(selected_tags)
                results.append({
                    **doc,
                    'score': score,
                    'matching_tags': matching_tags,
                    'context': [
                        f"Document tagged with: {', '.join(matching_tags)}"
                    ]
                })
        
        return sorted(
            results, key=lambda x: x['score'], reverse=True
        )[:top_k]
    
    def _keyword_search(
        self,
        query: str,
        top_k: int
    ) -> List[Dict[str, Any]]:
        """Simple keyword-based search"""
        results = []
        query_words = query.lower().split()
        
        for doc in self.documents:
            content_lower = doc['content'].lower()
            score = 0
            matches = []
            
            for word in query_words:
                count = content_lower.count(word)
                score += count
                if count > 0:
                    matches.append(word)
            
            if score > 0:
                context = self._get_context(doc['content'], query_words)
                results.append({
                    **doc,
                    'score': score,
                    'matches': matches,
                    'context': context
                })
        
        return sorted(
            results, key=lambda x: x['score'], reverse=True
        )[:top_k]
    
    def _similarity_search(
        self,
        query: str,
        top_k: int
    ) -> List[Dict[str, Any]]:
        """TF-IDF based similarity search"""
        if self.vectorizer is None:
            return self._keyword_search(query, top_k)
            
        query_vector = self.vectorizer.transform([query])
        similarities = cosine_similarity(
            query_vector, self.tfidf_matrix
        ).flatten()
        
        results = []
        for i, doc in enumerate(self.documents):
            if similarities[i] > 0:
                context = self._get_context(doc['content'], query.split())
                results.append({
                    **doc,
                    'score': similarities[i],
                    'context': context
                })
        
        return sorted(
            results, key=lambda x: x['score'], reverse=True
        )[:top_k]
    
    def _combined_search(
        self,
        query: str,
        top_k: int
    ) -> List[Dict[str, Any]]:
        """Combine keyword and similarity search"""
        keyword_results = self._keyword_search(query, top_k * 2)
        similarity_results = self._similarity_search(query, top_k * 2)
        
        # Combine and deduplicate
        combined = {}
        for result in keyword_results + similarity_results:
            filename = result['filename']
            if filename in combined:
                combined[filename]['score'] += result['score']
            else:
                combined[filename] = result
        
        return sorted(
            combined.values(), key=lambda x: x['score'], reverse=True
        )[:top_k]
    
    def _get_context(
        self,
        content: str,
        query_words: List[str],
        context_length: int = CONTEXT_LENGTH
    ) -> List[str]:
        """Extract context around query matches"""
        content_lower = content.lower()
        contexts = []
        
        for word in query_words:
            word = word.lower()
            start = content_lower.find(word)
            if start != -1:
                context_start = max(0, start - context_length // 2)
                context_end = min(
                    len(content), start + len(word) + context_length // 2
                )
                context = content[context_start:context_end]
                if context_start > 0:
                    context = "..." + context
                if context_end < len(content):
                    context = context + "..."
                contexts.append(context)
        
        return contexts[:MAX_CONTEXTS]

    def _categorize_documents(self) -> None:
        """Categorize and tag all documents"""
        if not self.documents:
            return
            
        categorized_docs = []
        for doc in self.documents:
            analysis = self.categorizer.analyze_document(
                doc['content'],
                doc['filename']
            )
            categorized_doc = {**doc, **analysis}
            categorized_docs.append(categorized_doc)
        
        self.documents = categorized_docs

    def get_category_summary(self) -> Dict[str, Any]:
        """Get summary statistics for categories and tags"""
        return self.categorizer.get_category_summary(self.documents)

    def filter_documents(
        self,
        category: Optional[str] = None,
        doc_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        urgency: Optional[str] = None,
        min_readability: Optional[float] = None,
        max_readability: Optional[float] = None
    ) -> List[Dict[str, Any]]:
        """Filter documents based on various criteria"""
        filtered = self.documents.copy()
        
        if category and category != 'All':
            filtered = [
                doc for doc in filtered
                if doc.get('business_category') == category
            ]
        
        if doc_type and doc_type != 'All':
            filtered = [
                doc for doc in filtered
                if doc.get('document_type') == doc_type
            ]
        
        if tags:
            for tag in tags:
                filtered = [
                    doc for doc in filtered
                    if tag in doc.get('tags', [])
                ]
        
        if urgency and urgency != 'All':
            filtered = [
                doc for doc in filtered
                if doc.get('urgency_level') == urgency
            ]
        
        if min_readability is not None:
            filtered = [
                doc for doc in filtered
                if doc.get('readability', 0) >= min_readability
            ]
        
        if max_readability is not None:
            filtered = [
                doc for doc in filtered
                if doc.get('readability', 100) <= max_readability
            ]
        
        return filtered

    def generate_wordcloud_tags(
        self,
        documents: Optional[List[Dict[str, Any]]] = None
    ) -> Optional[Tuple[Any, Counter]]:
        """Generate word cloud for document tags"""
        if documents is None:
            documents = self.documents
        
        if not documents:
            return None
            
        # Collect all tags with frequencies
        all_tags = []
        for doc in documents:
            all_tags.extend(doc.get('tags', []))
        
        if not all_tags:
            return None
            
        tag_freq = Counter(all_tags)
        
        # Generate word cloud
        wordcloud = WordCloud(**WORDCLOUD_CONFIG).generate_from_frequencies(
            tag_freq
        )
        
        return wordcloud, tag_freq
